from typing import Dict, Any, Optional
from open_notebook.domain.notebook import Note, Source
import json
import io

# For DOCX
from docx import Document
from docx.shared import Inches

# For PDF (Markdown -> HTML -> PDF)
import markdown # For converting Markdown to HTML
from weasyprint import HTML # For converting HTML to PDF

def note_to_txt(note: Note) -> str:
    """Converts a Note object to a plain text string."""
    return f"Title: {note.title or 'Untitled Note'}\n\n{note.content or ''}"

def note_to_md(note: Note) -> str:
    """Converts a Note object to a Markdown string."""
    return f"# {note.title or 'Untitled Note'}\n\n{note.content or ''}"

def note_to_json(note: Note) -> str:
    """Serializes a Note object to a JSON string."""
    note_data = {
        "id": note.id,
        "title": note.title or 'Untitled Note',
        "content": note.content or '',
        "note_type": note.note_type,
        "created": note.created.isoformat() if note.created else None,
        "updated": note.updated.isoformat() if note.updated else None,
    }
    return json.dumps(note_data, indent=2)

def note_to_docx_bytes(note: Note) -> bytes:
    """Converts a Note object to DOCX format (bytes)."""
    document = Document()
    document.add_heading(note.title or 'Untitled Note', level=1)
    
    # Add content. Handles None content gracefully.
    content = note.content or ""
    # Simple paragraph addition. For complex markdown, a more sophisticated parser might be needed.
    # For now, we treat the content as plain text for DOCX.
    # If content is Markdown, it will appear as raw Markdown text.
    document.add_paragraph(content)
    
    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()

def note_to_pdf_bytes(note: Note) -> bytes:
    """Converts a Note object to PDF bytes."""
    html_content = markdown.markdown(note.content or '')
    # Basic styling
    html_string = f"""
    <html>
        <head>
            <meta charset="UTF-8">
            <title>{note.title or 'Untitled Note'}</title>
            <style>
                body {{ font-family: sans-serif; margin: 2em; }}
                h1 {{ color: #333; }}
                p {{ line-height: 1.6; }}
            </style>
        </head>
        <body>
            <h1>{note.title or 'Untitled Note'}</h1>
            {html_content}
        </body>
    </html>
    """
    pdf_file = io.BytesIO()
    HTML(string=html_string).write_pdf(pdf_file)
    pdf_file.seek(0)
    return pdf_file.getvalue()

# --- Source Download Utilities ---

def source_to_txt(source: Source) -> str:
    """Converts a Source object to a plain text string."""
    content = source.full_text or ""
    if not content and source.insights:
        content = "\n\n".join([f"Insight ({i.insight_type}): \n{i.content}" for i in source.insights])
    elif not content and source.asset and source.asset.url:
        content = f"Source URL: {source.asset.url}"
    
    return f"Title: {source.title or 'Untitled Source'}\n\n{content}"

def source_to_md(source: Source) -> str:
    """Converts a Source object to a Markdown string."""
    content = source.full_text or ""
    if not content and source.insights:
        # Basic markdown for insights
        insight_md_parts = []
        for i in source.insights:
            insight_md_parts.append(f"## Insight: {i.insight_type}\n\n{i.content}")
        content = "\n\n---\n\n".join(insight_md_parts)
    elif not content and source.asset and source.asset.url:
        content = f"Source URL: [{source.asset.url}]({source.asset.url})"

    return f"# {source.title or 'Untitled Source'}\n\n{content}"

def source_to_json(source: Source) -> str:
    """Serializes a Source object to a JSON string."""
    source_data = {
        "id": source.id,
        "title": source.title,
        "full_text": source.full_text,
        "asset": source.asset.model_dump() if source.asset else None,
        "topics": source.topics,
        "insights": [insight.model_dump() for insight in source.insights] if source.insights else []
        # Add other relevant fields if necessary
    }
    return json.dumps(source_data, indent=4)

def source_to_docx_bytes(source: Source) -> bytes:
    """Converts a Source object to DOCX bytes (Word document)."""
    document = Document()
    document.add_heading(source.title or 'Untitled Source', level=1)
    
    content_to_add = source.full_text
    if not content_to_add and source.insights:
        for insight in source.insights:
            document.add_heading(f"Insight: {insight.insight_type}", level=2)
            document.add_paragraph(insight.content or '')
            document.add_paragraph() # Add some space
    elif not content_to_add and source.asset and source.asset.url:
        content_to_add = f"Source URL: {source.asset.url}"
        document.add_paragraph(content_to_add or '')
    else:
        document.add_paragraph(content_to_add or '')

    # Save to a BytesIO stream
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def source_to_pdf_bytes(source: Source) -> bytes:
    """Converts a Source object to PDF bytes."""
    # Consolidate content for PDF, similar to source_to_md but without the top-level title
    pdf_main_content = source.full_text or ""
    if not pdf_main_content and source.insights:
        insight_md_parts = []
        for i in source.insights:
            insight_md_parts.append(f"<h2>Insight: {i.insight_type}</h2>\\n\\n{i.content}") # Use h2 for insights in PDF
        pdf_main_content = "\\n\\n<hr/>\\n\\n".join(insight_md_parts)
    elif not pdf_main_content and source.asset and source.asset.url:
        # For PDF, make the URL clickable and clearly labeled
        pdf_main_content = f'<p>Source URL: <a href="{source.asset.url}">{source.asset.url}</a></p>'

    # Convert the consolidated content to HTML
    # If pdf_main_content already contains HTML (like the URL case or insights with h2),
    # markdown.markdown() should mostly pass it through.
    # If it's plain text or Markdown from source.full_text, it will be converted.
    html_body_content = markdown.markdown(pdf_main_content)
    
    html_string = f"""
    <html>
        <head>
            <meta charset="UTF-8">
            <title>{source.title or 'Untitled Source'}</title>
            <style>
                body {{ font-family: sans-serif; margin: 2em; }}
                h1 {{ color: #333; }} /* Main title */
                h2 {{ color: #555; margin-top: 1.5em; border-bottom: 1px solid #eee; padding-bottom: 0.3em;}} /* Insight titles */
                p, li {{ line-height: 1.6; }}
                hr {{ margin-top: 1em; margin-bottom: 1em; border: 0; border-top: 1px solid #eee; }}
                a {{ color: #0066cc; }}
            </style>
        </head>
        <body>
            <h1>{source.title or 'Untitled Source'}</h1>
            {html_body_content}
        </body>
    </html>
    """
    pdf_file = io.BytesIO()
    HTML(string=html_string).write_pdf(pdf_file)
    pdf_file.seek(0)
    return pdf_file.getvalue()

# --- Transformation Result Download Utilities ---

def transformation_to_txt(title: str, content: str) -> str:
    """Converts transformation result to a plain text string."""
    return f"Title: {title or 'Untitled Transformation'}\n\n{content or ''}"

def transformation_to_md(title: str, content: str) -> str:
    """Converts transformation result to a Markdown string."""
    return f"# {title or 'Untitled Transformation'}\n\n{content or ''}"

def transformation_to_json(title: str, content: str, original_source_id: Optional[str] = None, transformation_details: Optional[Dict[str, Any]] = None) -> str:
    """Serializes transformation result to a JSON string."""
    data = {
        "title": title,
        "content": content,
        "original_source_id": original_source_id,
        "transformation_details": transformation_details
    }
    return json.dumps(data, indent=4)

def transformation_to_docx_bytes(title: str, content: str) -> bytes:
    """Converts transformation result to DOCX bytes."""
    document = Document()
    document.add_heading(title or 'Untitled Transformation', level=1)
    document.add_paragraph(content or '')
    
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()

def transformation_to_pdf_bytes(title: str, content: str) -> bytes:
    """Converts transformation result to PDF bytes."""
    html_content = markdown.markdown(content or '')
    html_string = f"""
    <html>
        <head>
            <meta charset="UTF-8">
            <title>{title or 'Untitled Transformation'}</title>
            <style>
                body {{ font-family: sans-serif; margin: 2em; }}
                h1 {{ color: #333; }}
                p {{ line-height: 1.6; }}
            </style>
        </head>
        <body>
            <h1>{title or 'Untitled Transformation'}</h1>
            {html_content}
        </body>
    </html>
    """
    pdf_file = io.BytesIO()
    HTML(string=html_string).write_pdf(pdf_file)
    pdf_file.seek(0)
    return pdf_file.getvalue()

# Make sure to import Source from open_notebook.domain.notebook at the top
# from open_notebook.domain.notebook import Note, Source # Ensure Source is imported 